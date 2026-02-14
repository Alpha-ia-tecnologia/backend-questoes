from docx import Document
from docx.shared import Cm
from app.schemas.question_schema import QuestionSchema, QuestionWithImageSchema
from app.utils.save_image import save_image
import os

class GenerateDocxService:
    @staticmethod
    def generate_docx(questions: list[QuestionSchema | QuestionWithImageSchema], file_name: str):
        # Garante que a pasta export existe
        os.makedirs("export", exist_ok=True)
        
        doc = Document()
        doc._body.clear_content()
        
        doc.add_heading("Questões educacionais", 0)
        try:
            for question in questions:
                # Suporta tanto objetos Pydantic quanto dicts
                if hasattr(question, 'model_dump'):
                    q = question.model_dump()
                elif hasattr(question, 'dict'):
                    q = question.dict()
                else:
                    q = question if isinstance(question, dict) else dict(question)
                
                doc.add_heading(f"Questão {q['question_number']}", 3)
                doc.add_paragraph(f"Habilidade: {q['id_skill']} - {q['skill']}")
                doc.add_paragraph(f"Nível de proficiência: {q['proficiency_level']} - {q['proficiency_description']}")
                doc.add_heading(f"{q['title']}", 4)
                
                # Verifica se tem imagem (base64 ou URL)
                image_path = None
                
                # Debug: mostrar quais campos de imagem estão presentes
                has_base64 = bool(q.get('image_base64'))
                has_url = bool(q.get('image_url'))
                print(f"📸 Questão {q['question_number']}: image_base64={has_base64}, image_url={has_url}")
                if has_url:
                    print(f"   URL: {q.get('image_url')[:100] if q.get('image_url') else 'None'}...")
                
                if q.get('image_base64'):
                    try:
                        image_path = save_image(q['image_base64'])
                        print(f"   ✅ Imagem salva de base64: {image_path}")
                    except Exception as img_error:
                        print(f"   ❌ Erro ao salvar imagem base64: {img_error}")
                elif q.get('image_url'):
                    try:
                        # Se tem image_url, tenta ler do arquivo estático
                        # URL pode ser completa (http://...) ou relativa (/static/...)
                        url = q['image_url']
                        if '/static/' in url:
                            # Extrai o caminho relativo após /static/
                            relative_path = url.split('/static/')[-1]
                            local_path = os.path.join('static', relative_path)
                            print(f"   Verificando arquivo: {local_path}")
                            if os.path.exists(local_path):
                                image_path = local_path
                                print(f"   ✅ Arquivo encontrado: {local_path}")
                            else:
                                print(f"   ❌ Arquivo não encontrado: {local_path}")
                        else:
                            print(f"   ⚠️ URL não contém '/static/': {url}")
                    except Exception as img_error:
                        print(f"   ❌ Erro ao processar image_url: {img_error}")
                else:
                    print(f"   ⚠️ Nenhuma imagem disponível")
                
                if image_path:
                    try:
                        doc.add_picture(image_path, width=Cm(12), height=Cm(8))
                        print(f"   ✅ Imagem adicionada ao DOCX")
                    except Exception as img_error:
                        print(f"   ❌ Erro ao adicionar imagem ao documento: {img_error}")
                
                doc.add_paragraph(f"{q['text']}")
                doc.add_paragraph(f"{q['source']}")
                doc.add_paragraph("")
                doc.add_paragraph(f"{q['question_statement']}")
                doc.add_paragraph("")
                
                for alternative in q['alternatives']:
                    # Suporta tanto objetos quanto dicts
                    if isinstance(alternative, dict):
                        letter = alternative['letter']
                        text = alternative['text']
                    else:
                        letter = alternative.letter
                        text = alternative.text
                    doc.add_paragraph(f"({letter}) {text}")

                doc.add_paragraph("")
                doc.add_paragraph(f"Resposta correta: {q['correct_answer']}")
                doc.add_paragraph(f"Explicação: {q['explanation_question']}")
                doc.add_paragraph("")
                doc.add_paragraph("---")
                
            path = f"export/{file_name}.docx"
            doc.save(path)
            print(f"✅ Documento salvo em: {path}")
        except Exception as e:
            print(f"❌ Erro ao gerar documento: {e}")
            raise Exception(f"Error generating document: {e}")
